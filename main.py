import docx

def get_elements_from_docx(file_path):
    paragraphs = []
    images = []
    doc = docx.Document(file_path)

    for element in doc.element.body:
        if isinstance(element, docx.oxml.text.paragraph.CT_P):
            paragraph = docx.text.paragraph.Paragraph(element, doc)
            if paragraph.text.strip():
                paragraphs.append(paragraph)
        elif isinstance(element, docx.shape.InlineShape):
            images.append(element)

    return paragraphs, images

def mark_content_with_labels(paragraphs, images):
    for paragraph in paragraphs:
        paragraph.text += "-[This is a paragraph]"

    for image in images:
        image._element.get_or_add_ln().get_or_add_solidFill().get_or_add_srgbClr().val = "FFFFFF00"
        image.text += "-[This is a image]"

if __name__ == "__main__":

    input_file_path = 'test.docx'
    output_file_path = 'output_file2.docx'

    paragraphs, images = get_elements_from_docx(input_file_path)
    mark_content_with_labels(paragraphs, images)

    doc = docx.Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph.text, style=paragraph.style)

    for image in images:
        run = doc.add_paragraph().add_run()
        run._r.append(image._element)

    doc.save(output_file_path)
